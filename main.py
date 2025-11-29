# # from fastapi import FastAPI, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from routes.dyslexia_routes import router as dyslexia_router
# from docx import Document
# #from routes import data
# from pydantic import BaseModel
# from typing import Optional
# from jiwer import wer
# import tempfile, os
# #from routes.dyslexia_routes import router as dyslexia_router


# app = FastAPI()
# app.include_router(dyslexia_router)
# #app.include_router(dyslexia_router)

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # Allows all origins - you can specify domains later
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )


# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )


# # ----- Reference Texts -----
# TEXTS = [
#     {"id": "t1", "lang": "si", "content": "අද අහස පැහැදිලි ය. ළමයා පොතක් කියවයි."},
#     {"id": "t2", "lang": "en", "content": "The sky is clear today. The child reads a book."}
# ]
 
# def read_docx(file_path: str):
#     doc = Document(file_path)
#     lines = []

#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if text:
#             lines.append(text)

#     return lines


# def get_text(text_id: str):
#     for t in TEXTS:
#         if t["id"] == text_id:
#             return t
#     return None

# # ----- Utility -----
# def compute_metrics(reference: str, transcript: str, duration: Optional[float]):
#     ref_words = reference.split()
#     hyp_words = transcript.split()
#     total_ref = len(ref_words)
#     total_spoken = len(hyp_words)
#     error_rate = wer(reference.lower(), transcript.lower())
#     accuracy = max(0.0, 1 - error_rate) * 100
#     correct_words = round(accuracy/100 * total_ref, 2)
#     words_per_sec = None
#     if duration:
#         words_per_sec = round(total_spoken / duration, 2)
#     return {
#         "reference": reference,
#         "transcript": transcript,
#         "total_words": total_spoken,
#         "correct_words": correct_words,
#         "accuracy_percent": round(accuracy, 2),
#         "wer": round(error_rate, 4),
#         "words_per_second": words_per_sec
#     }


# # app.include_router(data.router)

# # @app.get("/")
# # def read_root():
# #     return {"message": "Welcome to RP Backend!"}

# # @app.get("/health")
# # def health_check():
# #     return {"status": "healthy", "service": "FastAPI Backend"}

# # ----- API -----
# @app.get("/texts")
# def list_texts():
#     return TEXTS

# class CompareBody(BaseModel):
#     text_id: str
#     transcript: str
#     duration: Optional[float] = None

# @app.post("/dyslexia/compare")
# def compare_text(body: CompareBody):
#     text = get_text(body.text_id)
#     if not text:
#         return {"error": "Invalid text id"}
#     metrics = compute_metrics(text["content"], body.transcript, body.duration)
#     return {"ok": True, "metrics": metrics}

# @app.post("/dyslexia/submit-audio")
# async def submit_audio(
#     text_id: str = Form(...),
#     duration: Optional[float] = Form(None),
#     file: UploadFile = File(...)
# ):
#     """Placeholder if you later enable transcription"""
#     text = get_text(text_id)
#     if not text:
#         return {"error": "Invalid text id"}

#     # Just save temporarily (no transcription yet)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
#         tmp.write(await file.read())
#         tmp_path = tmp.name

#     # For now, return mock result
#     try:
#         metrics = compute_metrics(text["content"], text["content"], duration)
#         return {"ok": True, "metrics": metrics, "note": "Audio received (transcription disabled)"}
#     finally:
#         os.remove(tmp_path)



from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from routes.dyslexia_routes import router as dyslexia_router

from docx import Document
from pydantic import BaseModel
from typing import Optional
from jiwer import wer
import tempfile
import os

# -----------------------------
app = FastAPI(
    title="Reading Proficiency (RP) Backend",
    description="API for dyslexia/reading assessment",
    version="1.0.0"
)

# Include your router (uncomment only one!)
app.include_router(dyslexia_router)  # This is correct if you want routes from routes/dyslexia_routes.py

# CORS - only one middleware is needed!
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],           # Change to specific origins in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------------
# Reference Texts
TEXTS = [
    {"id": "t1", "lang": "si", "content": "අද අහස පැහැදිලි ය. ළමයා පොතක් කියවයි."},
    {"id": "t2", "lang": "en", "content": "The sky is clear today. The child reads a book."}
]

def read_docx(file_path: str):
    doc = Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def get_text(text_id: str):
    for t in TEXTS:
        if t["id"] == text_id:
            return t
    return None

# -----------------------------
def compute_metrics(reference: str, transcript: str, duration: Optional[float] = None):
    error_rate = wer(reference.lower(), transcript.lower())
    accuracy = max(0.0, 1 - error_rate) * 100
    ref_words = reference.split()
    correct_words = round(accuracy / 100 * len(ref_words), 2)
    words_per_sec = round(len(transcript.split()) / duration, 2) if duration and duration > 0 else None

    return {
        "reference": reference,
        "transcript": transcript,
        "total_words": len(transcript.split()),
        "correct_words": correct_words,
        "accuracy_percent": round(accuracy, 2),
        "wer": round(error_rate, 4),
        "words_per_second": words_per_sec
    }

# -----------------------------
# Root endpoints
@app.get("/")
def read_root():
    return {"message": "Welcome to Reading Proficiency Backend!", "docs": "/docs"}

@app.get("/health")
def health_check():
    return {"status": "healthy"}

# -----------------------------
@app.get("/texts")
def list_texts():
    return {"texts": TEXTS}

# -----------------------------
class CompareBody(BaseModel):
    text_id: str
    transcript: str
    duration: Optional[float] = None

@app.post("/dyslexia/compare")
def compare_text(body: CompareBody):
    text = get_text(body.text_id)
    if not text:
        return {"ok": False, "error": "Invalid text_id"}
    
    metrics = compute_metrics(text["content"], body.transcript, body.duration)
    return {"ok": True, "metrics": metrics}

# -----------------------------
@app.post("/dyslexia/submit-audio")
async def submit_audio(
    text_id: str = Form(...),
    duration: Optional[float] = Form(None),
    file: UploadFile = File(...)
):
    text = get_text(text_id)
    if not text:
        return {"ok": False, "error": "Invalid text_id"}

    # Save uploaded audio temporarily
    suffix = os.path.splitext(file.filename)[1] or ".wav"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        content = await file.read()
        tmp_file.write(content)
        tmp_path = tmp_file.name

    try:
        # Placeholder: here you will later run speech-to-text
        # For now, just pretend the transcript is perfect (or mock it)
        mock_transcript = text["content"]  # Remove this line when real STT is added
        
        metrics = compute_metrics(text["content"], mock_transcript, duration)
        return {
            "ok": True,
            "metrics": metrics,
            "note": "Audio saved. Using mock transcript (STT not enabled yet).",
            "saved_file": tmp_path
        }
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)